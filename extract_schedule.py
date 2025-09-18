import sys
import os
import re
import subprocess
import tempfile
from docx import Document
import logging

# Настройка логирования
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

def convert_doc_to_docx(doc_path, temp_dir):
    """
    Конвертирует .doc файл в .docx с помощью libreoffice.
    Возвращает путь к .docx файлу или None в случае ошибки.
    """
    try:
        # Создаём временный файл для .docx
        temp_docx_path = os.path.join(temp_dir, os.path.basename(doc_path).replace('.doc', '.docx'))
        logging.info(f"Конвертируем {doc_path} в {temp_docx_path} с помощью libreoffice")

        # Проверяем, существует ли команда libreoffice
        result = subprocess.run(['which', 'libreoffice'], capture_output=True, text=True)
        if result.return_code != 0:
            logging.error(f"libreoffice не найден в системе: {result.stderr}")
            return None
        logging.debug(f"libreoffice найден: {result.stdout.strip()}")

        # Проверяем права доступа
        if not os.access(temp_dir, os.W_OK):
            logging.error(f"Нет прав на запись в {temp_dir}")
            return None
        if not os.access(doc_path, os.R_OK):
            logging.error(f"Нет прав на чтение {doc_path}")
            return None

        # Проверяем версию libreoffice
        version_result = subprocess.run(['libreoffice', '--version'], capture_output=True, text=True)
        logging.debug(f"Версия libreoffice: {version_result.stdout.strip()}")
        if version_result.return_code != 0:
            logging.error(f"Ошибка при проверке версии libreoffice: {version_result.stderr}")
            return None

        # Проверяем работоспособность libreoffice
        test_result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'txt', '/app/README.md', '--outdir', temp_dir],
            capture_output=True, text=True, timeout=30
        )
        logging.debug(f"Тестовая конверсия README.md: stdout={test_result.stdout}, stderr={test_result.stderr}, return_code={test_result.return_code}")

        # Запускаем libreoffice для конверсии
        logging.debug(f"Выполняем команду: libreoffice --headless --convert-to docx {doc_path} --outdir {temp_dir}")
        try:
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx', doc_path, '--outdir', temp_dir],
                capture_output=True, text=True, timeout=60
            )
            logging.debug(f"Команда libreoffice: {result.args}")
            logging.debug(f"Вывод libreoffice (stdout): {result.stdout}")
            logging.debug(f"Ошибки libreoffice (stderr): {result.stderr}")
            logging.debug(f"Код возврата libreoffice: {result.return_code}")

            if result.return_code != 0:
                logging.error(f"Ошибка конверсии {doc_path} в .docx: {result.stderr}")
                return None

            if not os.path.exists(temp_docx_path):
                logging.error(f"Файл {temp_docx_path} не был создан")
                return None

            logging.info(f"Успешно сконвертирован {doc_path} в {temp_docx_path}")
            return temp_docx_path
        except Exception as e:
            logging.error(f"Исключение при выполнении libreoffice: {type(e).__name__}: {str(e)}")
            return None
    except subprocess.TimeoutExpired:
        logging.error(f"Таймаут при конверсии {doc_path}")
        return None
    except Exception as e:
        logging.error(f"Ошибка при конверсии {doc_path} в .docx: {type(e).__name__}: {str(e)}")
        return None

def extract_doc_to_txt(doc_path, txt_path):
    """
    Извлекает текст из Word файла (.doc или .docx) и сохраняет в TXT файл,
    сохраняя структуру таблиц и обрабатывая отступы.
    """
    try:
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"Файл {doc_path} не найден")
        if not (doc_path.endswith('.doc') or doc_path.endswith('.docx')):
            raise ValueError("Входной файл должен иметь расширение .doc или .docx")

        logging.info(f"Обработка файла: {doc_path}")
        text_lines = []

        # Если файл .doc, конвертируем в .docx
        temp_docx_path = doc_path
        temp_dir = None
        if doc_path.endswith('.doc'):
            temp_dir = tempfile.mkdtemp()
            temp_docx_path = convert_doc_to_docx(doc_path, temp_dir)
            if not temp_docx_path:
                raise RuntimeError(f"Не удалось конвертировать {doc_path} в .docx")

        # Извлечение текста из .docx
        try:
            doc = Document(temp_docx_path)
            logging.info(f"Открыт файл: {temp_docx_path}")
            # Извлекаем текст из параграфов
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_lines.append(text)
                    logging.debug(f"Извлечен параграф: {text}")

            # Извлекаем текст из таблиц
            for table in doc.tables:
                max_columns = max(len(row.cells) for row in table.rows)
                logging.debug(f"Обработка таблицы с {max_columns} столбцами")
                for row in table.rows:
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    # Обработка объединённых ячеек
                    merged_cells = []
                    col_idx = 0
                    for cell in row.cells:
                        if cell._tc.grid_span > 1 or cell._tc.vMerge:
                            merged_cells.append(cell.text.strip().replace('\n', ' '))
                            for _ in range(cell._tc.grid_span - 1):
                                merged_cells.append('')
                                col_idx += 1
                        else:
                            merged_cells.append(cell.text.strip().replace('\n', ' '))
                        col_idx += 1
                    while len(merged_cells) < max_columns:
                        merged_cells.append('')
                    row_text = '│' + '│'.join(merged_cells) + '│'
                    text_lines.append(row_text)
                    logging.debug(f"Извлечена строка таблицы: {row_text}")
                text_lines.append('')
        except Exception as e:
            logging.error(f"Ошибка при обработке файла {temp_docx_path}: {e}")
            raise
        finally:
            if temp_dir:
                if os.path.exists(temp_docx_path):
                    os.remove(temp_docx_path)
                os.rmdir(temp_dir)

        if not text_lines or all(not line.strip() for line in text_lines):
            logging.warning(f"Файл {doc_path} пуст или не содержит полезного текста")
            return False

        # Сохраняем в TXT файл
        os.makedirs(os.path.dirname(txt_path), exist_ok=True)
        with open(txt_path, 'w', encoding='utf-8') as txt_file:
            for line in text_lines:
                txt_file.write(line + '\n')
        logging.info(f"Текст успешно извлечён из {doc_path} и сохранён в {txt_path}")

        # Проверяем содержимое файла
        if os.path.exists(txt_path):
            with open(txt_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                logging.info(f"Содержимое файла {txt_path} ({len(lines)} строк):")
                for i, line in enumerate(lines, 1):
                    logging.debug(f"Строка {i}: {line.strip()}")
            return True
        else:
            logging.error(f"Файл {txt_path} не был создан")
            return False

    except Exception as e:
        logging.error(f"Ошибка при обработке {doc_path}: {e}")
        return False

def main():
    """Основная функция для обработки всех .doc файлов в downloaded_schedules."""
    downloaded_dir = "downloaded_schedules"
    extracted_dir = "extracted_schedules"
    os.makedirs(extracted_dir, exist_ok=True)

    logging.info(f"Сканируем папку {downloaded_dir} на наличие .doc файлов")
    doc_files = [f for f in os.listdir(downloaded_dir) if f.endswith(('.doc', '.docx'))]
    logging.info(f"Найдено {len(doc_files)} файлов: {doc_files}")

    # Соответствие файлов дням недели
    day_mapping = {
        "rasp_monday.doc": "rasp_monday.txt",
        "rasp_tuesday.doc": "rasp_tuesday.txt",
        "rasp_wednesday.doc": "rasp_wednesday.txt",
        "rasp_thursday.doc": "rasp_thursday.txt",
        "rasp_friday.doc": "rasp_friday.txt",
        "rasp_saturday.doc": "rasp_saturday.txt",
    }

    success_count = 0
    error_count = 0

    for doc_file in doc_files:
        doc_path = os.path.join(downloaded_dir, doc_file)
        txt_file = day_mapping.get(doc_file, doc_file.replace('.docx', '.txt').replace('.doc', '.txt'))
        txt_path = os.path.join(extracted_dir, txt_file)
        logging.info(f"Обрабатываем {doc_path} -> {txt_path}")
        if extract_doc_to_txt(doc_path, txt_path):
            success_count += 1
        else:
            error_count += 1

    logging.info(f"Обработка завершена: {success_count} успешно, {error_count} ошибок")
    if error_count > 0:
        sys.exit(1)

if __name__ == "__main__":
    main()
